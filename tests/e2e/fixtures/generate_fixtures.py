"""Generate small test fixture files used by E2E document parse tests.

Creates the following files next to this script:
  sample.pdf   — 2-page PDF with text paragraphs and a simple data table
  sample.docx  — DOCX with headings, paragraphs, and a table
  sample.csv   — CSV with header row + 8 data rows
  sample.xlsx  — XLSX (single sheet) with header row + 8 data rows

Run directly::

    python tests/e2e/fixtures/generate_fixtures.py

Or import the helper functions for use in pytest fixtures::

    from tests.e2e.fixtures.generate_fixtures import generate_all
"""

from __future__ import annotations

from pathlib import Path

FIXTURES_DIR = Path(__file__).parent


# ── PDF ───────────────────────────────────────────────────────────────────────


def generate_pdf(dest: Path | None = None) -> Path:
    """Create a 2-page PDF with text and a table section."""
    import fitz  # pymupdf

    out = dest or FIXTURES_DIR / "sample.pdf"
    doc = fitz.open()

    # ── page 1: title + abstract ──────────────────────────────────────────────
    page1 = doc.new_page(width=595, height=842)  # A4
    page1.insert_text(
        (50, 80),
        "MantisFetch E2E Test Document",
        fontsize=18,
        fontname="helv",
    )
    page1.insert_text(
        (50, 120),
        "Abstract",
        fontsize=14,
        fontname="helv",
    )
    page1.insert_text(
        (50, 145),
        (
            "This document is a programmatically generated fixture used\n"
            "by the MantisFetch E2E test suite to verify PDF parsing,\n"
            "section extraction, and digest generation.\n"
            "It contains two pages with structured content."
        ),
        fontsize=11,
        fontname="helv",
    )
    page1.insert_text(
        (50, 260),
        "Introduction",
        fontsize=14,
        fontname="helv",
    )
    page1.insert_text(
        (50, 285),
        (
            "MantisFetch is an open-source data collection and document\n"
            "parsing platform. It supports PDF, DOCX, XLSX, and CSV\n"
            "formats with OCR fallback for scanned documents.\n\n"
            "The three-tier loading strategy (digest / brief / section)\n"
            "minimises token usage while preserving full content access."
        ),
        fontsize=11,
        fontname="helv",
    )

    # ── page 2: data section + table ─────────────────────────────────────────
    page2 = doc.new_page(width=595, height=842)
    page2.insert_text(
        (50, 80),
        "Performance Metrics",
        fontsize=14,
        fontname="helv",
    )
    page2.insert_text(
        (50, 105),
        (
            "The table below shows sample performance metrics collected\n"
            "during benchmark testing of the document parsing pipeline."
        ),
        fontsize=11,
        fontname="helv",
    )

    # Draw a simple table manually
    col_x = [50, 180, 310, 420]
    row_y = [150, 175, 200, 225, 250, 275]
    headers = ["Format", "Pages", "Time (s)", "Sections"]
    rows = [
        ["PDF", "45", "23.5", "12"],
        ["DOCX", "30", "8.2", "9"],
        ["XLSX", "—", "1.1", "3"],
        ["CSV", "—", "0.3", "1"],
    ]

    # Header row background (light grey via rect)
    page2.draw_rect(fitz.Rect(50, 140, 500, 162), color=(0.8, 0.8, 0.8), fill=(0.9, 0.9, 0.9))

    for i, h in enumerate(headers):
        page2.insert_text((col_x[i] + 4, 157), h, fontsize=10, fontname="helv")

    for r_idx, row in enumerate(rows):
        y = row_y[r_idx + 1]
        for c_idx, cell in enumerate(row):
            page2.insert_text((col_x[c_idx] + 4, y - 3), cell, fontsize=10, fontname="helv")

    # Table border
    page2.draw_rect(fitz.Rect(50, 140, 500, 285), color=(0.3, 0.3, 0.3), width=0.5)

    page2.insert_text(
        (50, 320),
        "Conclusion",
        fontsize=14,
        fontname="helv",
    )
    page2.insert_text(
        (50, 345),
        (
            "MantisFetch delivers efficient document processing with\n"
            "minimal resource usage. The provider abstraction layer\n"
            "enables seamless switching between LLM backends."
        ),
        fontsize=11,
        fontname="helv",
    )

    doc.save(str(out))
    doc.close()
    print(f"Generated: {out}")
    return out


# ── DOCX ──────────────────────────────────────────────────────────────────────


def generate_docx(dest: Path | None = None) -> Path:
    """Create a DOCX with headings, paragraphs, and a table."""
    from docx import Document
    from docx.shared import Pt

    out = dest or FIXTURES_DIR / "sample.docx"
    doc = Document()

    doc.add_heading("MantisFetch E2E Test Document", level=1)

    doc.add_heading("Abstract", level=2)
    doc.add_paragraph(
        "This document is a programmatically generated fixture used "
        "by the MantisFetch E2E test suite to verify DOCX parsing, "
        "section extraction, and digest generation."
    )

    doc.add_heading("Introduction", level=2)
    doc.add_paragraph(
        "MantisFetch is an open-source data collection and document "
        "parsing platform by ReadyForAI. It provides a unified API "
        "for web capture and document parsing with LLM-powered summaries."
    )
    doc.add_paragraph(
        "The system supports PDF, DOCX, XLSX, and CSV formats. "
        "OCR fallback is available for scanned or image-heavy PDFs. "
        "A three-tier loading strategy minimises token costs."
    )

    doc.add_heading("Performance Metrics", level=2)
    doc.add_paragraph(
        "The following table summarises benchmark results across file formats."
    )

    # Table
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    headers = ["Format", "Pages", "Time (s)", "Sections"]
    rows = [
        ("PDF", "45", "23.5", "12"),
        ("DOCX", "30", "8.2", "9"),
        ("XLSX", "—", "1.1", "3"),
        ("CSV", "—", "0.3", "1"),
    ]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = val

    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph(
        "MantisFetch delivers efficient document processing with minimal "
        "resource usage. The multi-LLM provider abstraction enables "
        "seamless switching between Gemini, OpenAI, Ollama, and other backends."
    )

    doc.save(str(out))
    print(f"Generated: {out}")
    return out


# ── CSV ───────────────────────────────────────────────────────────────────────


def generate_csv(dest: Path | None = None) -> Path:
    """Create a CSV with a header row and 8 data rows."""
    import csv

    out = dest or FIXTURES_DIR / "sample.csv"
    rows = [
        ["Document", "Format", "Pages", "Sections", "Tables", "Time_s", "OCR_pages"],
        ["annual_report.pdf", "pdf", "120", "18", "6", "58.2", "3"],
        ["q3_results.pdf", "pdf", "45", "12", "8", "23.5", "0"],
        ["contract_draft.docx", "docx", "30", "9", "2", "8.2", "0"],
        ["meeting_notes.docx", "docx", "8", "5", "0", "2.1", "0"],
        ["revenue_data.xlsx", "xlsx", "1", "3", "3", "1.1", "0"],
        ["headcount.xlsx", "xlsx", "1", "2", "2", "0.9", "0"],
        ["survey_results.csv", "csv", "1", "1", "1", "0.3", "0"],
        ["expense_report.csv", "csv", "1", "1", "1", "0.4", "0"],
    ]
    with out.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerows(rows)
    print(f"Generated: {out}")
    return out


# ── XLSX ──────────────────────────────────────────────────────────────────────


def generate_xlsx(dest: Path | None = None) -> Path:
    """Create an XLSX with one sheet, a header row, and 8 data rows."""
    import openpyxl
    from openpyxl.styles import Font

    out = dest or FIXTURES_DIR / "sample.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "ParseMetrics"

    headers = ["Document", "Format", "Pages", "Sections", "Tables", "Time_s", "OCR_pages"]
    data = [
        ["annual_report.pdf", "pdf", 120, 18, 6, 58.2, 3],
        ["q3_results.pdf", "pdf", 45, 12, 8, 23.5, 0],
        ["contract_draft.docx", "docx", 30, 9, 2, 8.2, 0],
        ["meeting_notes.docx", "docx", 8, 5, 0, 2.1, 0],
        ["revenue_data.xlsx", "xlsx", 1, 3, 3, 1.1, 0],
        ["headcount.xlsx", "xlsx", 1, 2, 2, 0.9, 0],
        ["survey_results.csv", "csv", 1, 1, 1, 0.3, 0],
        ["expense_report.csv", "csv", 1, 1, 1, 0.4, 0],
    ]

    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)

    for row in data:
        ws.append(row)

    wb.save(str(out))
    print(f"Generated: {out}")
    return out


# ── orchestrator ──────────────────────────────────────────────────────────────


def generate_all(dest_dir: Path | None = None) -> dict[str, Path]:
    """Generate all fixture files and return a dict mapping format → path."""
    d = dest_dir or FIXTURES_DIR
    d.mkdir(parents=True, exist_ok=True)
    return {
        "pdf": generate_pdf(d / "sample.pdf"),
        "docx": generate_docx(d / "sample.docx"),
        "csv": generate_csv(d / "sample.csv"),
        "xlsx": generate_xlsx(d / "sample.xlsx"),
    }


if __name__ == "__main__":
    paths = generate_all()
    print("\nAll fixtures generated:")
    for fmt, path in paths.items():
        print(f"  {fmt}: {path} ({path.stat().st_size} bytes)")
